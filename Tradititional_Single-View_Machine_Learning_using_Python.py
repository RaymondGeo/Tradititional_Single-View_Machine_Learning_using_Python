# Import required libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, GridSearchCV, cross_val_score
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error
from sklearn.neighbors import KNeighborsRegressor
from sklearn.ensemble import RandomForestRegressor
from sklearn.svm import SVR
from sklearn.neural_network import MLPRegressor
import xgboost as xgb
from sklearn.inspection import permutation_importance
import warnings
warnings.filterwarnings('ignore')

import os
from datetime import datetime
import openpyxl
from docx import Document
from docx.shared import Inches
import logging

# Set random seed for reproducibility
np.random.seed(123)

class MLPipeline:
    def __init__(self, base_output_path):
        """Initialize the ML Pipeline"""
        self.base_output_path = base_output_path
        self.train_proportions = [0.7, 0.8]
        self.normalization_params = {}
        self.all_results = []
        
        #"For optimal performance, it is recommended to extend the hyperparameter search space based on the dataset characteristics and available computational resources"
        # Define algorithms and their hyperparameter spaces 
        self.algorithms = {
            'knn': {
                'name': 'K-Nearest Neighbors',
                'model': KNeighborsRegressor(),
                'params': {
                    'n_neighbors': [1, 3, 5, 7, 10, 15, 20]  
                }
            },
            'rf': {
                'name': 'Random Forest',
                'model': RandomForestRegressor(random_state=123),
                'params': {
                    'n_estimators': [50, 100, 200, 300, 500],  
                    'max_features': None,  # Will be set dynamically (reduced range)
                    'min_samples_leaf': [1, 2, 5, 10]  
                }
            },
            'svm': {
                'name': 'Support Vector Machine',
                'model': SVR(),
                'params': [
                    {
                        'kernel': ['linear'],
                        'C': [0.1, 0.5, 1.0, 2.0, 5.0, 10.0]  
                    },
                    {
                        'kernel': ['rbf'],
                        'C': [0.1, 0.5, 1.0, 2.0, 5.0, 10.0],  
                        'gamma': [0.1, 0.5, 1.0]  #
                    }
                ]
            },
            'xgboost': {
                'name': 'XGBoost',
                'model': xgb.XGBRegressor(random_state=123, verbosity=0),
                'params': {
                    'n_estimators': [50, 100, 200, 300],  
                    'learning_rate': [0.01, 0.05, 0.1, 0.2, 0.3],  
                    'max_depth': [2, 4, 6, 8], 
                    'subsample': [0.7, 0.8, 0.9, 1.0],  
                    'colsample_bytree': [0.7, 0.8, 0.9, 1.0]  
                }
            },
            'nnet': {
                'name': 'Neural Network',
                'model': MLPRegressor(random_state=123),
                'params': {
                    'hidden_layer_sizes': [(5,), (10,), (15,), (20,)],  
                    'alpha': [0.0, 0.01, 0.05, 0.1],  
                    'max_iter': [200, 300, 500] 
                }
            }
        }
    
    def calculate_composite_score(self, r2, rmse, r2_values, rmse_values, r2_weight=0.5, rmse_weight=0.5):
        """
        Calculate composite score considering both R² (higher is better) and RMSE (lower is better)
        
        Parameters:
        - r2: R² value for current model
        - rmse: RMSE value for current model  
        - r2_values: All R² values for normalization
        - rmse_values: All RMSE values for normalization
        - r2_weight: Weight for R² component (default 0.5)
        - rmse_weight: Weight for RMSE component (default 0.5)
        """
        # Normalize R² (0-1 scale, higher is better)
        r2_min, r2_max = min(r2_values), max(r2_values)
        if r2_max == r2_min:
            r2_normalized = 1.0
        else:
            r2_normalized = (r2 - r2_min) / (r2_max - r2_min)
        
        # Normalize RMSE and invert (0-1 scale, higher is better after inversion)
        rmse_min, rmse_max = min(rmse_values), max(rmse_values)
        if rmse_max == rmse_min:
            rmse_normalized_inverted = 1.0
        else:
            rmse_normalized = (rmse - rmse_min) / (rmse_max - rmse_min)
            rmse_normalized_inverted = 1 - rmse_normalized
        
        # Calculate composite score
        composite_score = (r2_weight * r2_normalized) + (rmse_weight * rmse_normalized_inverted)
        
        return composite_score, r2_normalized, rmse_normalized_inverted
    
    def load_data(self, file_path):
        """Load data from Excel file"""
        try:
            data = pd.read_excel(file_path)
            print(f"Data loaded successfully. Shape: {data.shape}")
            return data
        except Exception as e:
            print(f"Error loading data: {e}")
            return None
    
    def normalize_data(self, train_data, test_data, target_col='species'):
        """Normalize features and target using training data statistics"""
        # Identify predictor columns
        predictor_cols = [col for col in train_data.columns if col != target_col]
        
        # Initialize scalers
        feature_scaler = StandardScaler()
        target_scaler = StandardScaler()
        
        # Create copies for normalization
        train_normalized = train_data.copy()
        test_normalized = test_data.copy()
        
        # Normalize features
        train_normalized[predictor_cols] = feature_scaler.fit_transform(train_data[predictor_cols])
        test_normalized[predictor_cols] = feature_scaler.transform(test_data[predictor_cols])
        
        # Normalize target
        train_normalized[target_col] = target_scaler.fit_transform(train_data[[target_col]]).ravel()
        test_normalized[target_col] = target_scaler.transform(test_data[[target_col]]).ravel()
        
        # Store normalization parameters
        self.normalization_params = {
            'feature_scaler': feature_scaler,
            'target_scaler': target_scaler,
            'predictor_cols': predictor_cols,
            'target_col': target_col
        }
        
        return train_normalized, test_normalized
    
    def denormalize_predictions(self, normalized_predictions):
        """Denormalize predictions back to original scale"""
        return self.normalization_params['target_scaler'].inverse_transform(
            normalized_predictions.reshape(-1, 1)
        ).ravel()
    
    def calculate_metrics(self, y_true, y_pred):
        """Calculate R², RMSE, and MAE"""
        r2 = r2_score(y_true, y_pred)
        rmse = np.sqrt(mean_squared_error(y_true, y_pred))
        mae = mean_absolute_error(y_true, y_pred)
        return {'r2': r2, 'rmse': rmse, 'mae': mae}
    
    def calculate_feature_importance(self, model, X_test, y_test):
        """Calculate permutation-based feature importance"""
        try:
            # Calculate permutation importance
            perm_importance = permutation_importance(
                model, X_test, y_test, 
                n_repeats=10, 
                random_state=123,
                scoring='neg_mean_squared_error'
            )
            
            # Convert to percentage
            importance_scores = perm_importance.importances_mean
            importance_scores = np.maximum(importance_scores, 0)  # Ensure non-negative
            
            if importance_scores.sum() == 0:
                importance_scores = np.ones(len(importance_scores))
            
            importance_percentages = (importance_scores / importance_scores.sum()) * 100
            
            # Create DataFrame
            feature_importance = pd.DataFrame({
                'Feature': self.normalization_params['predictor_cols'],
                'Importance': importance_percentages
            }).sort_values('Importance', ascending=False)
            
            return feature_importance
        except Exception as e:
            print(f"Error calculating feature importance: {e}")
            return pd.DataFrame({'Feature': [], 'Importance': []})
    
    def format_params(self, params):
        """Format hyperparameters for display"""
        if not params:
            return "No hyperparameters tuned."
        return ", ".join([f"{k}={v}" for k, v in params.items()])
    
    def create_plots(self, algorithm_name, y_true, y_pred, feature_importance, output_dir):
        """Create visualization plots"""
        plt.style.use('default')
        
        # 1. Actual vs Predicted plot
        plt.figure(figsize=(8, 6))
        plt.scatter(y_true, y_pred, alpha=0.6)
        plt.plot([y_true.min(), y_true.max()], [y_true.min(), y_true.max()], 'r--', lw=2)
        plt.xlabel('Actual species Score')
        plt.ylabel('Predicted species Score')
        plt.title(f'{algorithm_name} - Actual vs Predicted species Score')
        plt.grid(True, alpha=0.3)
        plot_path = os.path.join(output_dir, f'{algorithm_name.lower().replace(" ", "_")}_plot.png')
        plt.savefig(plot_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        # 2. Feature importance plot
        if not feature_importance.empty and len(feature_importance) > 0:
            plt.figure(figsize=(10, 8))
            top_features = feature_importance.head(15)  # Show top 15 features
            plt.barh(range(len(top_features)), top_features['Importance'])
            plt.yticks(range(len(top_features)), top_features['Feature'])
            plt.xlabel('Importance (%)')
            plt.title(f'{algorithm_name} - Feature Importance')
            plt.gca().invert_yaxis()
            
            # Add percentage labels
            for i, v in enumerate(top_features['Importance']):
                plt.text(v + 0.1, i, f'{v:.1f}%', va='center')
            
            plt.tight_layout()
            importance_plot_path = os.path.join(output_dir, f'{algorithm_name.lower().replace(" ", "_")}_feature_importance.png')
            plt.savefig(importance_plot_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            return plot_path, importance_plot_path
        
        return plot_path, None
    
    def create_algorithm_report(self, algorithm_name, best_params, train_metrics, test_metrics, 
                              feature_importance, output_dir):
        """Create individual algorithm report"""
        doc = Document()
        doc.add_heading(f'{algorithm_name} Report', 0)
        
        # Model description
        doc.add_heading('Model Description', level=1)
        doc.add_paragraph(f'Algorithm: {algorithm_name}')
        doc.add_paragraph(f'Hyperparameters: {self.format_params(best_params)}')
        doc.add_paragraph('Note: Data was normalized before training and predictions were denormalized for evaluation.')
        
        # Performance metrics
        doc.add_heading('Performance Metrics', level=1)
        
        # Create metrics table
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Training'
        hdr_cells[2].text = 'Testing'
        
        # Data rows
        metrics_data = [
            ('R-squared', f"{train_metrics['r2']:.4f}", f"{test_metrics['r2']:.4f}"),
            ('RMSE', f"{train_metrics['rmse']:.4f}", f"{test_metrics['rmse']:.4f}"),
            ('MAE', f"{train_metrics['mae']:.4f}", f"{test_metrics['mae']:.4f}")
        ]
        
        for i, (metric, train_val, test_val) in enumerate(metrics_data, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = metric
            row_cells[1].text = train_val
            row_cells[2].text = test_val
        
        # Feature importance section
        if not feature_importance.empty:
            doc.add_heading('Feature Importance', level=1)
            doc.add_paragraph(
                'Feature importance was calculated using a permutation-based approach. '
                'This method measures how much the model\'s performance decreases when each feature is randomly shuffled.'
            )
            
            # Add top 10 features table
            top_features = feature_importance.head(10)
            importance_table = doc.add_table(rows=len(top_features) + 1, cols=2)
            importance_table.style = 'Table Grid'
            
            # Headers
            hdr_cells = importance_table.rows[0].cells
            hdr_cells[0].text = 'Feature'
            hdr_cells[1].text = 'Importance (%)'
            
            # Data
            for i, (_, row) in enumerate(top_features.iterrows(), 1):
                row_cells = importance_table.rows[i].cells
                row_cells[0].text = str(row['Feature'])
                row_cells[1].text = f"{row['Importance']:.2f}"
        
        # Save report
        report_path = os.path.join(output_dir, f'{algorithm_name.lower().replace(" ", "_")}_report.docx')
        doc.save(report_path)
        
        return report_path
    
    def setup_rf_params(self, num_features):
        """Set up Random Forest parameters dynamically based on number of features (REDUCED for speed)"""
        import math
        max_mtry = max(1, int(math.sqrt(num_features)))
        
        # Create reduced mtry values: include 1, sqrt(features), and max
        if max_mtry <= 3:
            mtry_values = list(range(1, max_mtry + 1))
        else:
            mtry_values = [1, 2, max_mtry//2, max_mtry]
        
        return {
            'n_estimators': [50, 100, 200, 300, 500],  # Reduced to 5 values
            'max_features': mtry_values,  # Reduced to key values only
            'min_samples_leaf': [1, 2, 5, 10]  # Reduced to 4 values
        }
    
    def run_algorithm(self, algorithm_key, train_data_norm, test_data_norm, 
                     train_data_orig, test_data_orig, output_dir):
        """Run a single algorithm with hyperparameter tuning"""
        algorithm = self.algorithms[algorithm_key]
        algorithm_name = algorithm['name']
        
        print(f"\nProcessing Algorithm: {algorithm_name}")
        
        try:
            # Prepare data
            X_train = train_data_norm[self.normalization_params['predictor_cols']]
            y_train = train_data_norm[self.normalization_params['target_col']]
            X_test = test_data_norm[self.normalization_params['predictor_cols']]
            y_test = test_data_norm[self.normalization_params['target_col']]
            
            # Handle Random Forest dynamic parameters
            if algorithm_key == 'rf':
                num_features = len(self.normalization_params['predictor_cols'])
                rf_params = self.setup_rf_params(num_features)
                print(f"Random Forest: Using mtry range 1 to {max(rf_params['max_features'])} (based on {num_features} features)")
            else:
                rf_params = algorithm['params']
            
            # Grid search with cross-validation
            grid_search = GridSearchCV(
                algorithm['model'],
                rf_params if algorithm_key == 'rf' else algorithm['params'],
                cv=5,
                scoring='neg_mean_squared_error',
                n_jobs=-1
            )
            
            print("Training model...")
            grid_search.fit(X_train, y_train)
            
            # Get best model and parameters
            best_model = grid_search.best_estimator_
            best_params = grid_search.best_params_
            
            # Make predictions (normalized)
            train_pred_norm = best_model.predict(X_train)
            test_pred_norm = best_model.predict(X_test)
            
            # Denormalize predictions
            train_pred = self.denormalize_predictions(train_pred_norm)
            test_pred = self.denormalize_predictions(test_pred_norm)
            
            # Calculate metrics on original scale
            train_metrics = self.calculate_metrics(train_data_orig[self.normalization_params['target_col']], train_pred)
            test_metrics = self.calculate_metrics(test_data_orig[self.normalization_params['target_col']], test_pred)
            
            # Calculate feature importance
            print("Calculating feature importance...")
            feature_importance = self.calculate_feature_importance(best_model, X_test, y_test)
            
            # Create visualizations
            plot_path, importance_plot_path = self.create_plots(
                algorithm_name, 
                test_data_orig[self.normalization_params['target_col']], 
                test_pred, 
                feature_importance, 
                output_dir
            )
            
            # Save feature importance to CSV
            if not feature_importance.empty:
                feature_importance.to_csv(
                    os.path.join(output_dir, f'{algorithm_key}_feature_importance.csv'), 
                    index=False
                )
            
            # Save predictions to Excel
            pred_data = pd.DataFrame({
                'Actual_species': test_data_orig[self.normalization_params['target_col']],
                'Predicted_species': test_pred
            })
            pred_data.to_excel(
                os.path.join(output_dir, f'{algorithm_key}_predictions.xlsx'), 
                index=False
            )
            
            # Create individual report
            self.create_algorithm_report(
                algorithm_name, best_params, train_metrics, test_metrics, 
                feature_importance, output_dir
            )
            
            print(f"Completed {algorithm_name}")
            
            return {
                'algorithm': algorithm_name,
                'train_metrics': train_metrics,
                'test_metrics': test_metrics,
                'best_params': best_params,
                'predictions': {'train': train_pred, 'test': test_pred},
                'feature_importance': feature_importance
            }
            
        except Exception as e:
            print(f"Error in {algorithm_name}: {e}")
            # Create error log
            with open(os.path.join(output_dir, f'{algorithm_key}_error_log.txt'), 'w') as f:
                f.write(f"Error in algorithm: {algorithm_name}\n")
                f.write(f"Error message: {str(e)}\n")
            return None
    
    def create_comparison_plots(self, results_df, output_dir, train_prop):
        """Create comparison plots across algorithms"""
        plt.style.use('default')
        
        # Calculate percentages with rounding to avoid floating point issues
        train_pct = round(train_prop * 100)
        test_pct = round((1 - train_prop) * 100)
        
        # Sort by composite score for plotting
        sorted_df = results_df.sort_values('Composite_Score', ascending=False)
        
        # Composite Score comparison
        plt.figure(figsize=(12, 6))
        bars = plt.bar(range(len(sorted_df)), sorted_df['Composite_Score'], color='purple', alpha=0.7)
        plt.xticks(range(len(sorted_df)), sorted_df['Algorithm'], rotation=45, ha='right')
        plt.ylabel('Composite Score (higher is better)')
        plt.title(f'Algorithm Performance Comparison (Composite Score: R² + RMSE) - {train_pct}-{test_pct} Split')
        
        # Add value labels on bars
        for i, bar in enumerate(bars):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                    f'{height:.3f}', ha='center', va='bottom')
        
        plt.tight_layout()
        composite_plot_path = os.path.join(output_dir, 'composite_score_comparison.png')
        plt.savefig(composite_plot_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        # Test RMSE comparison
        plt.figure(figsize=(10, 6))
        sorted_df_rmse = results_df.sort_values('Test_RMSE')
        bars = plt.bar(range(len(sorted_df_rmse)), sorted_df_rmse['Test_RMSE'])
        plt.xticks(range(len(sorted_df_rmse)), sorted_df_rmse['Algorithm'], rotation=45, ha='right')
        plt.ylabel('Test RMSE (lower is better)')
        plt.title(f'Algorithm Performance Comparison (Test RMSE) - {train_pct}-{test_pct} Split')
        
        # Add value labels on bars
        for i, bar in enumerate(bars):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.001,
                    f'{height:.4f}', ha='center', va='bottom')
        
        plt.tight_layout()
        test_plot_path = os.path.join(output_dir, 'test_rmse_comparison.png')
        plt.savefig(test_plot_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        # Test R² comparison
        plt.figure(figsize=(10, 6))
        sorted_df_r2 = results_df.sort_values('Test_R2', ascending=False)
        bars = plt.bar(range(len(sorted_df_r2)), sorted_df_r2['Test_R2'], color='green', alpha=0.7)
        plt.xticks(range(len(sorted_df_r2)), sorted_df_r2['Algorithm'], rotation=45, ha='right')
        plt.ylabel('Test R² (higher is better)')
        plt.title(f'Algorithm Performance Comparison (Test R²) - {train_pct}-{test_pct} Split')
        
        # Add value labels on bars
        for i, bar in enumerate(bars):
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                    f'{height:.4f}', ha='center', va='bottom')
        
        plt.tight_layout()
        r2_plot_path = os.path.join(output_dir, 'test_r2_comparison.png')
        plt.savefig(r2_plot_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return composite_plot_path, test_plot_path, r2_plot_path
    
    def create_feature_importance_heatmap(self, all_feature_importance, output_dir):
        """Create feature importance heatmap across algorithms"""
        if not all_feature_importance:
            return None
            
        # Combine all feature importance data
        combined_importance = pd.DataFrame()
        for alg_name, importance_df in all_feature_importance.items():
            if not importance_df.empty:
                temp_df = importance_df.copy()
                temp_df['Algorithm'] = alg_name
                combined_importance = pd.concat([combined_importance, temp_df], ignore_index=True)
        
        if combined_importance.empty:
            return None
        
        # Calculate average importance
        avg_importance = combined_importance.groupby('Feature')['Importance'].mean().sort_values(ascending=False)
        
        # Get top 10 features
        top_features = avg_importance.head(10).index.tolist()
        
        # Create pivot table for heatmap
        heatmap_data = combined_importance[combined_importance['Feature'].isin(top_features)]
        pivot_data = heatmap_data.pivot(index='Feature', columns='Algorithm', values='Importance')
        pivot_data = pivot_data.reindex(top_features)  # Maintain order by importance
        
        # Create heatmap
        plt.figure(figsize=(12, 8))
        sns.heatmap(pivot_data, annot=True, fmt='.1f', cmap='Blues', cbar_kws={'label': 'Importance (%)'})
        plt.title('Top 10 Features Importance (%) Across Algorithms')
        plt.xlabel('Algorithm')
        plt.ylabel('Feature')
        plt.tight_layout()
        
        heatmap_path = os.path.join(output_dir, 'feature_importance_heatmap.png')
        plt.savefig(heatmap_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        # Save average importance
        avg_importance_df = pd.DataFrame({
            'Feature': avg_importance.index,
            'Importance': avg_importance.values
        })
        avg_importance_df.to_csv(os.path.join(output_dir, 'average_feature_importance.csv'), index=False)
        
        # Create average importance plot
        plt.figure(figsize=(10, 8))
        top_10_avg = avg_importance_df.head(10)
        bars = plt.barh(range(len(top_10_avg)), top_10_avg['Importance'], color='darkgreen')
        plt.yticks(range(len(top_10_avg)), top_10_avg['Feature'])
        plt.xlabel('Importance (%)')
        plt.title('Average Feature Importance Across All Models')
        plt.gca().invert_yaxis()
        
        # Add percentage labels
        for i, bar in enumerate(bars):
            width = bar.get_width()
            plt.text(width + 0.1, bar.get_y() + bar.get_height()/2.,
                    f'{width:.1f}%', ha='left', va='center')
        
        plt.tight_layout()
        avg_plot_path = os.path.join(output_dir, 'average_feature_importance.png')
        plt.savefig(avg_plot_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return heatmap_path, avg_plot_path, avg_importance_df
    
    def create_summary_report(self, results_df, all_feature_importance, output_dir, train_prop):
        """Create comprehensive summary report"""
        doc = Document()
        train_pct = round(train_prop * 100)
        test_pct = round((1 - train_prop) * 100)
        prop_string = f"{train_pct}-{test_pct}"
        doc.add_heading(f'species Score Prediction - Algorithm Comparison with Composite Scoring ({prop_string} Split)', 0)
        
        # Experiment overview
        doc.add_heading('Experiment Overview', level=1)
        doc.add_paragraph(
            f'This document summarizes the results of {len(results_df)} machine learning algorithms '
            f'for predicting species Score. Each algorithm was trained on normalized data and tuned using 5-fold cross-validation '
            f'with a focus on minimizing RMSE (Root Mean Square Error). Predictions were denormalized before evaluation on a held-out test set.\n\n'
            f'Data was split into {train_pct}% training and {test_pct}% testing.\n\n'
            f'**Hyperparameter Optimization:** Parameter grids were optimized to focus on the most effective ranges, '
            f'balancing comprehensive search with computational efficiency.\n\n'
            f'**NEW: Composite Scoring Method**\n'
            f'The best model is selected using a composite score that considers both R² (higher is better) and RMSE (lower is better). '
            f'Both metrics are normalized to 0-1 scale and combined with equal weights (50% each) to create a comprehensive performance measure.'
        )
        
        # Performance comparison
        doc.add_heading('Algorithm Performance Comparison (Sorted by Composite Score)', level=1)
        
        # Create table
        table = doc.add_table(rows=len(results_df) + 1, cols=6)
        table.style = 'Table Grid'
        
        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Algorithm'
        hdr_cells[1].text = 'Composite Score'
        hdr_cells[2].text = 'Test R²'
        hdr_cells[3].text = 'Test RMSE'
        hdr_cells[4].text = 'Test MAE'
        hdr_cells[5].text = 'Rank'
        
        # Data
        sorted_results = results_df.sort_values('Composite_Score', ascending=False)
        for i, (_, row) in enumerate(sorted_results.iterrows(), 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(row['Algorithm'])
            row_cells[1].text = f"{row['Composite_Score']:.4f}"
            row_cells[2].text = f"{row['Test_R2']:.4f}"
            row_cells[3].text = f"{row['Test_RMSE']:.4f}"
            row_cells[4].text = f"{row['Test_MAE']:.4f}"
            row_cells[5].text = str(i)
        
        # Feature importance analysis
        if all_feature_importance:
            doc.add_heading('Feature Importance Analysis', level=1)
            doc.add_paragraph(
                'Feature importance was calculated using a model-agnostic permutation-based approach. '
                'This method measures how much model performance decreases when each feature is randomly shuffled. '
                'The importance values are normalized to percentages that sum to 100% for each model, '
                'making it easy to interpret the relative contribution of each feature.'
            )
        
        # Best algorithm recommendation
        if not sorted_results.empty:
            best_alg = sorted_results.iloc[0]
            doc.add_heading('Recommendation', level=1)
            doc.add_paragraph(
                f'The algorithm with the highest Composite Score (considering both R² and RMSE) was {best_alg["Algorithm"]} with:\n'
                f'Composite Score: {best_alg["Composite_Score"]:.4f}\n'
                f'Test R²: {best_alg["Test_R2"]:.4f}\n'
                f'Test RMSE: {best_alg["Test_RMSE"]:.4f}\n'
                f'Test MAE: {best_alg["Test_MAE"]:.4f}\n\n'
                f'This selection balances both predictive accuracy (R²) and prediction precision (RMSE), '
                f'providing the most comprehensive model performance.'
            )
        
        # Save report
        report_path = os.path.join(output_dir, 'final_summary_report.docx')
        doc.save(report_path)
        
        return report_path
    
    def run_ml_pipeline(self, train_prop, data_file_path):
        """Run the entire ML pipeline for a given training proportion"""
        train_pct = round(train_prop * 100)
        test_pct = round((1 - train_prop) * 100)
        
        print(f"\n{'='*50}")
        print(f"Running with training proportion: {train_pct}%, testing proportion: {test_pct}%")
        print(f"{'='*50}")
        
        # Create output directory (using already calculated percentages)
        prop_string = f"{train_pct}-{test_pct}"
        output_dir = os.path.join(self.base_output_path, f"ML_Results_{prop_string}")
        os.makedirs(output_dir, exist_ok=True)
        
        # Load data
        data = self.load_data(data_file_path)
        if data is None:
            return None
        
        # Split data
        train_data, test_data = train_test_split(
            data, test_size=1-train_prop, random_state=123, stratify=None
        )
        
        print(f"Training set size: {len(train_data)}, Testing set size: {len(test_data)}")
        
        # Normalize data
        train_data_norm, test_data_norm = self.normalize_data(train_data, test_data)
        
        # Save normalization parameters
        norm_params_df = pd.DataFrame({
            'Variable': ['Feature_Mean', 'Feature_Std', 'Target_Mean', 'Target_Std'],
            'Values': [
                str(self.normalization_params['feature_scaler'].mean_),
                str(self.normalization_params['feature_scaler'].scale_),
                self.normalization_params['target_scaler'].mean_[0],
                self.normalization_params['target_scaler'].scale_[0]
            ]
        })
        norm_params_df.to_csv(os.path.join(output_dir, 'normalization_parameters.csv'), index=False)
        
        # Initialize results
        results_list = []
        all_predictions = {}
        all_feature_importance = {}
        
        # Run each algorithm
        for alg_key in self.algorithms.keys():
            result = self.run_algorithm(
                alg_key, train_data_norm, test_data_norm, 
                train_data, test_data, output_dir
            )
            
            if result:
                results_list.append({
                    'Algorithm': result['algorithm'],
                    'Train_R2': result['train_metrics']['r2'],
                    'Train_RMSE': result['train_metrics']['rmse'],
                    'Train_MAE': result['train_metrics']['mae'],
                    'Test_R2': result['test_metrics']['r2'],
                    'Test_RMSE': result['test_metrics']['rmse'],
                    'Test_MAE': result['test_metrics']['mae'],
                    'Parameters': self.format_params(result['best_params'])
                })
                
                all_predictions[alg_key] = result['predictions']
                all_feature_importance[alg_key] = result['feature_importance']
        
        # Create results DataFrame
        results_df = pd.DataFrame(results_list)
        
        if not results_df.empty:
            # Calculate composite scores
            r2_values = results_df['Test_R2'].values
            rmse_values = results_df['Test_RMSE'].values
            
            composite_scores = []
            r2_normalized_list = []
            rmse_normalized_list = []
            
            for _, row in results_df.iterrows():
                comp_score, r2_norm, rmse_norm = self.calculate_composite_score(
                    row['Test_R2'], row['Test_RMSE'], r2_values, rmse_values
                )
                composite_scores.append(comp_score)
                r2_normalized_list.append(r2_norm)
                rmse_normalized_list.append(rmse_norm)
            
            results_df['Composite_Score'] = composite_scores
            results_df['R2_Normalized'] = r2_normalized_list
            results_df['RMSE_Normalized_Inverted'] = rmse_normalized_list
            
            # Save results sorted by composite score
            results_df_sorted = results_df.sort_values('Composite_Score', ascending=False)
            results_df_sorted.to_csv(os.path.join(output_dir, 'algorithm_comparison_by_composite_score.csv'), index=False)
            
            # Also save results sorted by RMSE for reference
            results_df_sorted_rmse = results_df.sort_values('Test_RMSE')
            results_df_sorted_rmse.to_csv(os.path.join(output_dir, 'algorithm_comparison_by_test_rmse.csv'), index=False)
            
            # Create comparison plots
            self.create_comparison_plots(results_df, output_dir, train_prop)
            
            # Create feature importance analysis
            if all_feature_importance:
                self.create_feature_importance_heatmap(all_feature_importance, output_dir)
            
            # Create consolidated predictions file
            if all_predictions:
                consolidated_pred = pd.DataFrame({
                    'Actual_species': test_data['species'].values
                })
                
                for alg_key, preds in all_predictions.items():
                    consolidated_pred[f'{self.algorithms[alg_key]["name"]}_Predicted'] = preds['test']
                
                consolidated_pred.to_excel(os.path.join(output_dir, 'all_predictions_comparison.xlsx'), index=False)
            
            # Create summary report
            self.create_summary_report(results_df, all_feature_importance, output_dir, train_prop)
            
            print(f"\nAnalysis completed for {prop_string} split. Results saved in: {output_dir}")
            print(f"Best algorithm by Composite Score: {results_df_sorted.iloc[0]['Algorithm']} (Score: {results_df_sorted.iloc[0]['Composite_Score']:.4f})")
            
            return {
                'train_prop': train_prop,
                'output_dir': output_dir,
                'best_algorithm': results_df_sorted.iloc[0]['Algorithm'] if not results_df_sorted.empty else None,
                'best_composite_score': results_df_sorted.iloc[0]['Composite_Score'] if not results_df_sorted.empty else None,
                'best_test_r2': results_df_sorted.iloc[0]['Test_R2'] if not results_df_sorted.empty else None,
                'best_test_rmse': results_df_sorted.iloc[0]['Test_RMSE'] if not results_df_sorted.empty else None,
                'results_df': results_df
            }
        
        return None
    
    def run_all_proportions(self, data_file_path):
        """Run ML pipeline for all training proportions"""
        print("Starting comprehensive ML analysis across multiple training proportions...")
        print("Using COMPOSITE SCORING method (R² + RMSE) for best model selection...")
        
        for train_prop in self.train_proportions:
            train_pct = round(train_prop * 100)
            print(f"\n{'='*60}")
            print(f"STARTING ANALYSIS WITH {train_pct}% TRAINING DATA")
            print(f"{'='*60}")
            
            result = self.run_ml_pipeline(train_prop, data_file_path)
            if result:
                self.all_results.append(result)
        
        # Create master comparison
        self.create_master_comparison()
        
        print(f"\n{'='*60}")
        print("COMPLETED ALL ANALYSES")
        print(f"{'='*60}")
    
    def create_master_comparison(self):
        """Create master comparison across all training proportions"""
        if not self.all_results:
            return
        
        # Create master directory
        master_dir = os.path.join(self.base_output_path, "ML_Results_Combined")
        os.makedirs(master_dir, exist_ok=True)
        
        # Create comparison DataFrame
        comparison_data = []
        for result in self.all_results:
            if result['best_algorithm'] and result['best_composite_score']:
                comparison_data.append({
                    'Training_Proportion': result['train_prop'],
                    'Testing_Proportion': 1 - result['train_prop'],
                    'Best_Algorithm': result['best_algorithm'],
                    'Best_Composite_Score': result['best_composite_score'],
                    'Best_Test_R2': result['best_test_r2'],
                    'Best_Test_RMSE': result['best_test_rmse']
                })
        
        comparison_df = pd.DataFrame(comparison_data)
        comparison_df.to_csv(os.path.join(master_dir, 'training_proportion_comparison_composite.csv'), index=False)
        
        # Create comparison plot
        if not comparison_df.empty:
            plt.figure(figsize=(12, 8))
            
            # Create subplot for composite score
            plt.subplot(2, 1, 1)
            bars = plt.bar(range(len(comparison_df)), comparison_df['Best_Composite_Score'])
            
            # Color bars by algorithm
            unique_algs = comparison_df['Best_Algorithm'].unique()
            colors = plt.cm.Set3(np.linspace(0, 1, len(unique_algs)))
            alg_color_map = dict(zip(unique_algs, colors))
            
            for i, (bar, alg) in enumerate(zip(bars, comparison_df['Best_Algorithm'])):
                bar.set_color(alg_color_map[alg])
                # Add value label
                plt.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.005,
                        f'{bar.get_height():.3f}', ha='center', va='bottom')
            
            plt.xticks(range(len(comparison_df)), 
                      [f'{round(x*100)}%' for x in comparison_df['Training_Proportion']])
            plt.xlabel('Training Proportion')
            plt.ylabel('Best Composite Score')
            plt.title('Best Composite Score by Training Proportion')
            plt.grid(True, alpha=0.3)
            
            # Create subplot for RMSE
            plt.subplot(2, 1, 2)
            bars2 = plt.bar(range(len(comparison_df)), comparison_df['Best_Test_RMSE'])
            
            for i, (bar, alg) in enumerate(zip(bars2, comparison_df['Best_Algorithm'])):
                bar.set_color(alg_color_map[alg])
                # Add value label
                plt.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.001,
                        f'{bar.get_height():.4f}', ha='center', va='bottom')
            
            plt.xticks(range(len(comparison_df)), 
                      [f'{round(x*100)}%' for x in comparison_df['Training_Proportion']])
            plt.xlabel('Training Proportion')
            plt.ylabel('Best Test RMSE')
            plt.title('Best Test RMSE by Training Proportion')
            plt.grid(True, alpha=0.3)
            
            # Create legend
            handles = [plt.Rectangle((0,0),1,1, color=alg_color_map[alg]) for alg in unique_algs]
            plt.figlegend(handles, unique_algs, title='Best Algorithm', loc='center right', bbox_to_anchor=(1.15, 0.5))
            
            plt.tight_layout()
            plt.savefig(os.path.join(master_dir, 'training_proportion_comparison_composite.png'), dpi=300, bbox_inches='tight')
            plt.close()
        
        # Create master summary document
        doc = Document()
        doc.add_heading('species Score Prediction - Training Proportion Comparison (Composite Scoring)', 0)
        
        doc.add_heading('Overview', level=1)
        doc.add_paragraph(
            'This document compares the performance of machine learning models for predicting species Score '
            'using different training/testing data splits. For each split, multiple algorithms were trained and evaluated '
            'with the same methodology, and the best model was selected using a COMPOSITE SCORING approach that '
            'considers both R² (predictive accuracy) and RMSE (prediction precision) with equal weights.'
        )
        
        doc.add_heading('Composite Scoring Method', level=1)
        doc.add_paragraph(
            'The composite score is calculated by:\n'
            '1. Normalizing R² values to 0-1 scale (higher R² = better score)\n'
            '2. Normalizing RMSE values to 0-1 scale and inverting (lower RMSE = better score)\n'
            '3. Combining both with equal weights (50% each): Composite Score = 0.5 × R²_normalized + 0.5 × (1 - RMSE_normalized)\n\n'
            'This ensures that the selected model has the best balance of both predictive accuracy and precision.'
        )
        
        if not comparison_df.empty:
            # Add comparison table
            doc.add_heading('Performance Across Training Proportions', level=1)
            table = doc.add_table(rows=len(comparison_df) + 1, cols=6)
            table.style = 'Table Grid'
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Training %'
            hdr_cells[1].text = 'Testing %'
            hdr_cells[2].text = 'Best Algorithm'
            hdr_cells[3].text = 'Composite Score'
            hdr_cells[4].text = 'Test R²'
            hdr_cells[5].text = 'Test RMSE'
            
            # Data
            for i, (_, row) in enumerate(comparison_df.iterrows(), 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = f"{round(row['Training_Proportion']*100)}%"
                row_cells[1].text = f"{round(row['Testing_Proportion']*100)}%"
                row_cells[2].text = str(row['Best_Algorithm'])
                row_cells[3].text = f"{row['Best_Composite_Score']:.4f}"
                row_cells[4].text = f"{row['Best_Test_R2']:.4f}"
                row_cells[5].text = f"{row['Best_Test_RMSE']:.4f}"
            
            # Best recommendation
            best_idx = comparison_df['Best_Composite_Score'].idxmax()
            best_result = comparison_df.loc[best_idx]
            
            doc.add_heading('Recommendation', level=1)
            doc.add_paragraph(
                f'Based on the composite scoring comparison across different training proportions, the optimal configuration is:\n\n'
                f'• Training proportion: {round(best_result["Training_Proportion"]*100)}%\n'
                f'• Algorithm: {best_result["Best_Algorithm"]}\n'
                f'• Composite Score: {best_result["Best_Composite_Score"]:.4f}\n'
                f'• Test R²: {best_result["Best_Test_R2"]:.4f}\n'
                f'• Test RMSE: {best_result["Best_Test_RMSE"]:.4f}\n\n'
                f'This configuration achieved the highest composite score, indicating the best balance between '
                f'predictive accuracy (R²) and prediction precision (RMSE).'
            )
        
        # Save master summary
        doc.save(os.path.join(master_dir, 'training_proportion_summary_composite.docx'))
        
        print(f"Master comparison with composite scoring saved in: {master_dir}")

# Example usage
if __name__ == "__main__":
    # Configure paths
    base_output_path = r"C:\Enter Your Base Path"
    data_file_path = r"C:\iris_dataset.xlsx"
    
    # Create and run pipeline
    pipeline = MLPipeline(base_output_path)
    pipeline.run_all_proportions(data_file_path)
    
    print("\nAll analyses completed successfully!")
    print(f"Results saved in: {base_output_path}")
    print("\n*** USING COMPOSITE SCORING: Best models selected based on HIGHEST R² AND LOWEST RMSE ***")
    print("*** OPTIMIZED HYPERPARAMETERS: Reduced grid sizes for faster execution while maintaining performance ***")